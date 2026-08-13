from __future__ import annotations

from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
MANUAL = ROOT / "docs" / "游戏说明书.docx"

REPLACEMENTS = {
    "抵达格子后，可选择执行该格子效果（购买、打工、打卡等）。若为事件格，须抽取并执行事件牌。":
        "抵达格子后，可选择执行该格子效果（购买、打工、打卡等）。若实际移动到事件格，进入行动阶段时自动公开抽取并执行事件牌；除牌面明确规定外，事件不会默认结束或跳过行动阶段，结算后重置完整15秒并继续行动。",
    "玩家仅可在卡牌标注的可用阶段打出非遗卡牌（每张限用一次，打出后归还至非遗点），效果详见卡牌说明。":
        "玩家仅可在卡牌标注的合法时机使用非遗卡牌。当前交易所尚未实装，已使用的主动型非遗牌暂时移出本局；国家级非遗牌正常收集与计分，但暂不能使用。",
    "2.可选择保留事件牌不公开，但每张牌仅能使用一次（含保留后使用）。":
        "2.仅妙手回春、游目骋怀、畅行无阻、金蝉脱壳、移花接木可保留；抽取时公开，保留后仅持有者可见。",
    "3.牌组用完后置入弃牌堆，不可重复使用。":
        "3.每张事件牌每局只结算一次，使用后进入弃牌堆且不洗回；牌库抽完时事件格无事发生。",
    "4.事件牌共40张。":
        "4.事件定义共40张；当前35张进入牌库。鉴往知来、釜底抽薪、展艺共研、市集淘珍等待交易所，孤注一掷等待职业技能。",
    "·神话传说：任意时刻消耗，无效化或转移他人使用的食物牌/事件牌效果。":
        "·神话传说：他人食物牌或事件牌效果即将作用于自己时消耗，无效化自己受到的部分，或将该部分转移给另一名合法玩家；转移后的新目标仍可继续响应。",
}


def replace_paragraph_text(paragraph, new_text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)


def main() -> None:
    document = Document(MANUAL)
    replaced = set()
    for paragraph in document.paragraphs:
        original = paragraph.text
        if original in REPLACEMENTS:
            replace_paragraph_text(paragraph, REPLACEMENTS[original])
            replaced.add(original)
    missing = set(REPLACEMENTS) - replaced
    if missing:
        raise RuntimeError(f"说明书中未找到待替换段落：{sorted(missing)}")

    # 复用事件规则后的空段落，补充计时与响应链，不改变后续章节位置。
    event_heading_index = next(i for i, p in enumerate(document.paragraphs) if p.text == "事件牌")
    blank_after_rules = document.paragraphs[event_heading_index + 5]
    if blank_after_rules.text:
        raise RuntimeError("事件规则后的预留段落不再为空，拒绝覆盖。")
    replace_paragraph_text(
        blank_after_rules,
        "5.事件选择与响应期间暂停行动计时，每一步有独立15秒；群体效果只处理响应者自己的部分，转移后的新目标仍可继续响应。",
    )
    document.save(MANUAL)
    print(MANUAL)


if __name__ == "__main__":
    main()
